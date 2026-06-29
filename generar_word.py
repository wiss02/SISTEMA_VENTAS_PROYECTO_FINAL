import os
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Por favor, instala python-docx ejecutando: pip install python-docx")
    exit()

def create_apa_document():
    doc = Document()

    # Configuración APA básica (Times New Roman 12)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # 1. Carátula
    for _ in range(4):
        doc.add_paragraph()
    
    title = doc.add_paragraph("Sistema de Gestión Inmobiliaria TEM742\nProyecto Final Monolítico")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_format = title.runs[0]
    title_format.bold = True
    
    doc.add_paragraph()
    
    info = [
        "Wilson Rene [Tus Apellidos]",
        "Materia: Tecnologías Emergentes II",
        "Proyecto de curso I-2026",
        "Docente: M. Sc. Mario Tórrez C.",
        "28 de junio de 2026"
    ]
    for line in info:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    doc.add_page_break()

    # 2. Planteamiento y Justificación
    h1 = doc.add_heading('Planteamiento del problema y justificación', level=1)
    doc.add_paragraph(
        "En el rubro de bienes raíces, la velocidad de respuesta y la organización de la información son factores críticos para el éxito. "
        "Actualmente, un gran porcentaje de agencias inmobiliarias tradicionales aún gestiona su catálogo de propiedades, la base de datos de propietarios y la agenda de visitas mediante métodos rudimentarios. "
        "Esta falta de sistematización ocasiona pérdida de datos sensibles, lentitud en la atención y una clara desventaja competitiva en un mercado cada vez más digitalizado."
    )
    doc.add_paragraph(
        "El desarrollo e implementación del Sistema de Gestión Inmobiliaria 'TEM742' nace como una solución directa a esta problemática. "
        "Justificamos la creación de esta plataforma web empresarial (monolítica) porque proporciona un ecosistema centralizado donde los agentes pueden administrar todo el flujo de trabajo en tiempo real. "
        "La digitalización de estos procesos no solo minimiza el margen de error humano, sino que eleva el perfil profesional de la agencia."
    )

    # 3. Objetivos
    doc.add_heading('Objetivo general y objetivos específicos', level=1)
    doc.add_heading('Objetivo General', level=2)
    doc.add_paragraph(
        "Diseñar, desarrollar e implementar una aplicación web empresarial bajo arquitectura monolítica para la gestión integral de operaciones de una agencia inmobiliaria, "
        "permitiendo la administración eficiente del catálogo de inmuebles, el control de usuarios, programación de visitas y seguimiento de contratos, garantizando la integridad de los datos mediante una base de datos relacional sólida."
    )
    doc.add_heading('Objetivos Específicos', level=2)
    doc.add_paragraph("1. Implementar un sistema de seguridad y autenticación estricto con control de acceso basado en roles (RBAC: Administrador, Agente, Cliente).")
    doc.add_paragraph("2. Desarrollar un catálogo digital interactivo de inmuebles que soporte operaciones CRUD.")
    doc.add_paragraph("3. Diseñar y estructurar una base de datos relacional normalizada de 8 tablas utilizando SQLAlchemy (ORM).")
    doc.add_paragraph("4. Construir un Panel Administrativo (Dashboard) que procese los datos en tiempo real y muestre métricas.")
    doc.add_paragraph("5. Aplicar el patrón de diseño Application Factory utilizando el framework Flask.")

    # 4. Módulos
    doc.add_page_break()
    doc.add_heading('Descripción de los módulos principales', level=1)
    
    modulos = [
        ("Módulo de Autenticación, Seguridad y Roles", "Se encarga de validar la identidad de quienes ingresan a la plataforma. Utiliza técnicas avanzadas de hashing de contraseñas. Adapta dinámicamente el menú lateral según el rol."),
        ("Módulo Core: Catálogo de Propiedades", "Muestra todas las propiedades disponibles. Permite filtrar dinámicamente por tipo, acción y buscar por texto libre."),
        ("Módulo de Propietarios e Interacción Rápida", "Permite registrar rápidamente a un nuevo propietario mediante un modal interactivo sin necesidad de abandonar el formulario actual."),
        ("Módulo Gerencial: Dashboard Analítico", "Proporciona tarjetas indicadoras con estadísticas vitales y gráficos visuales dinámicos.")
    ]

    for titulo, desc in modulos:
        doc.add_heading(titulo, level=2)
        doc.add_paragraph(desc)
        p_img = doc.add_paragraph("[==========================================================]\n"
                                  "[                 PEGAR CAPTURA DE PANTALLA AQUÍ           ]\n"
                                  "[==========================================================]")
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.runs[0].bold = True

    # 5. Diagrama E-R
    doc.add_page_break()
    doc.add_heading('Modelo de base de datos (diagrama relacional)', level=1)
    doc.add_paragraph("Para soportar toda la lógica empresarial, se diseñó un modelo relacional estrictamente normalizado de 8 tablas interconectadas que eliminan la redundancia y mantienen la coherencia referencial.")
    p_img2 = doc.add_paragraph("[==========================================================]\n"
                               "[      PEGAR IMAGEN DEL DIAGRAMA GENERADO EN DBDIAGRAM     ]\n"
                               "[==========================================================]")
    p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img2.runs[0].bold = True

    # 6. Tecnologías
    doc.add_heading('Descripción de Tecnologías utilizadas', level=1)
    doc.add_paragraph("- Backend: Python 3.13 y Flask (Application Factory).")
    doc.add_paragraph("- Base de Datos: SQLite (para desarrollo) integrado con el ORM Flask-SQLAlchemy.")
    doc.add_paragraph("- Frontend: HTML5, CSS3, JavaScript y Bootstrap 5 (Responsive Design).")
    doc.add_paragraph("- Seguridad: Werkzeug.security y Flask-Login.")

    # 7. Herramientas IA
    doc.add_heading('Uso de herramientas de Inteligencia Artificial', level=1)
    doc.add_paragraph(
        "Se utilizaron Modelos de Lenguaje Grande (LLMs) como asistencia técnica. "
        "Las consultas se enfocaron principalmente en estructurar la arquitectura del software bajo el patrón Application Factory y resolución de conflictos de Claves Foráneas (Foreign Keys). "
        "La Inteligencia Artificial fungió como un mentor técnico, permitiendo al equipo concentrarse en la lógica de negocio y acelerando el tiempo de desarrollo."
    )

    # 8. Conclusiones
    doc.add_heading('Conclusiones y aprendizajes', level=1)
    doc.add_heading('Conclusiones', level=2)
    doc.add_paragraph(
        "El proyecto 'Sistema de Gestión Inmobiliaria TEM742' fue culminado con éxito, cumpliendo rigurosamente con los requerimientos técnicos. "
        "Se demostró la viabilidad de una arquitectura monolítica. La aplicación es completamente funcional, segura y estructurada profesionalmente."
    )
    doc.add_heading('Aprendizajes', level=2)
    doc.add_paragraph("- Dominio del patrón de diseño Application Factory y Blueprints en Flask.")
    doc.add_paragraph("- Uso práctico de un ORM (SQLAlchemy) para independizar el código del motor de base de datos.")
    doc.add_paragraph("- Integración de interfaces responsivas y usables mediante Bootstrap 5.")
    doc.add_paragraph("- Afianzamiento de conocimientos en ciberseguridad básica (Hasheo y Roles).")

    doc.save('INFORME_FINAL_TEM742.docx')
    print("Documento guardado como 'INFORME_FINAL_TEM742.docx'")

if __name__ == '__main__':
    create_apa_document()
