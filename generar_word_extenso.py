from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_extended_apa_document():
    print("Generando el documento Word EXTENSO con imágenes incrustadas...")
    doc = Document()

    # Configuración APA 7 estricta (Times New Roman 12, Doble espacio, Sangría 0.5")
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Formato de párrafo para APA 7
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 2.0
    paragraph_format.first_line_indent = Inches(0.5)

    # 1. Carátula
    for _ in range(5):
        doc.add_paragraph()
    
    title = doc.add_paragraph("Sistema de Gestión Inmobiliaria TEM742\nProyecto Final de Arquitectura Monolítica\n\n")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_format = title.runs[0]
    title_format.bold = True
    
    doc.add_paragraph()
    
    info = [
        "Presentado por: Wilson Rene",
        "Materia: Tecnologías Emergentes II",
        "Proyecto de curso I-2026",
        "Docente: M. Sc. Mario Tórrez C.",
        "28 de junio de 2026"
    ]
    for line in info:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    doc.add_page_break()

    # 2. Introducción y Marco Teórico
    doc.add_heading('1. Introducción y Marco Teórico', level=1)
    doc.add_paragraph(
        "El presente documento detalla el análisis, diseño e implementación del Sistema de Gestión Inmobiliaria 'TEM742'. "
        "En la era digital contemporánea, la gestión de la información representa el activo más valioso para cualquier organización, "
        "y el sector de bienes raíces no es la excepción. La transición de métodos analógicos a plataformas digitales integradas "
        "ha demostrado ser un factor determinante en la competitividad de las agencias inmobiliarias a nivel global."
    )
    doc.add_paragraph(
        "Desde una perspectiva técnica, este proyecto se fundamenta en el paradigma de Arquitectura Monolítica. "
        "A diferencia de los microservicios, un monolito consolida la interfaz de usuario, la lógica de negocio y el acceso a datos "
        "en una única base de código desplegable. Esta elección arquitectónica se justifica por la escala inicial del proyecto, "
        "permitiendo un desarrollo ágil, pruebas unificadas y despliegues simplificados, manteniendo al mismo tiempo "
        "una modularidad interna estricta mediante el uso de Blueprints en el ecosistema Flask."
    )
    doc.add_paragraph(
        "Adicionalmente, se emplea el modelo de Control de Acceso Basado en Roles (RBAC, por sus siglas en inglés). "
        "Este modelo de seguridad informática restringe el acceso al sistema basándose en los roles que los usuarios "
        "desempeñan dentro de la organización, garantizando el principio de mínimo privilegio."
    )

    doc.add_page_break()

    # 3. Planteamiento y Justificación
    doc.add_heading('2. Planteamiento del problema y justificación', level=1)
    doc.add_heading('2.1. Planteamiento del Problema', level=2)
    doc.add_paragraph(
        "En el rubro de bienes raíces, la velocidad de respuesta y la organización de la información son factores críticos para el cierre exitoso de ventas. "
        "Actualmente, un gran porcentaje de agencias inmobiliarias tradicionales aún gestiona su catálogo de propiedades, "
        "la vasta base de datos de propietarios y la compleja agenda de visitas mediante métodos rudimentarios como hojas de cálculo aisladas o registros manuales. "
        "Esta carencia de sistematización tecnológica ocasiona múltiples cuellos de botella: pérdida de datos sensibles, "
        "lentitud extrema en la atención a prospectos, duplicidad de esfuerzos en el seguimiento de contratos y una clara desventaja competitiva."
    )
    doc.add_heading('2.2. Justificación', level=2)
    doc.add_paragraph(
        "El desarrollo e implementación del Sistema 'TEM742' nace como una solución de ingeniería de software directa a esta problemática. "
        "Justificamos la creación de esta plataforma web empresarial porque proporciona un ecosistema centralizado donde los agentes pueden "
        "administrar todo el flujo de trabajo operativo en tiempo real. La digitalización de estos procesos no solo minimiza drásticamente "
        "el margen de error humano (por ejemplo, evitar agendar dos visitas simultáneas para la misma propiedad), sino que eleva el perfil profesional de la agencia frente a sus clientes."
    )
    
    doc.add_page_break()

    # 4. Objetivos
    doc.add_heading('3. Objetivo general y objetivos específicos', level=1)
    doc.add_heading('3.1. Objetivo General', level=2)
    doc.add_paragraph(
        "Diseñar, desarrollar e implementar una aplicación web empresarial bajo arquitectura monolítica para la gestión integral de operaciones de una agencia inmobiliaria, "
        "permitiendo la administración eficiente del catálogo de inmuebles, el control estricto de usuarios, la programación automatizada de visitas y el seguimiento del ciclo de vida de los contratos, "
        "garantizando la integridad relacional de los datos y proporcionando una interfaz de usuario altamente interactiva."
    )
    doc.add_heading('3.2. Objetivos Específicos', level=2)
    doc.add_paragraph("1. Implementar un módulo de seguridad y autenticación criptográfica con control de acceso basado en roles (Administrador, Agente, Cliente).")
    doc.add_paragraph("2. Desarrollar un catálogo digital interactivo de inmuebles que soporte operaciones CRUD completas y persistencia de archivos multimedia.")
    doc.add_paragraph("3. Diseñar y estructurar una base de datos relacional altamente normalizada (Tercera Forma Normal) compuesta por 8 entidades principales utilizando SQLAlchemy.")
    doc.add_paragraph("4. Construir un Panel de Control Gerencial (Dashboard) que procese datos agregados en tiempo real y exponga métricas de rendimiento a través de visualizaciones gráficas.")
    doc.add_paragraph("5. Aplicar el patrón de diseño Application Factory y Blueprints para garantizar la mantenibilidad a largo plazo del código fuente.")
    doc.add_paragraph("6. Implementar técnicas de diseño web responsivo (Mobile First) para asegurar la compatibilidad multiplataforma de la aplicación.")

    doc.add_page_break()

    # 5. Módulos y Capturas
    doc.add_heading('4. Descripción detallada de los módulos principales', level=1)
    
    modulos = [
        ("4.1. Módulo de Autenticación, Seguridad y Roles (RBAC)", 
         "Este módulo representa la barrera de seguridad perimetral del sistema. Se encarga de validar la identidad de los usuarios mediante el algoritmo de hash de contraseñas de Werkzeug. "
         "Al autenticarse, el sistema inyecta el contexto del rol en la sesión, adaptando dinámicamente el menú de navegación y restringiendo el acceso a endpoints protegidos. "
         "Un cliente estándar jamás podrá visualizar la lógica de facturación de un agente.", 
         "capturas/login.png"),
        
        ("4.2. Módulo Core: Catálogo de Propiedades y Gestión de Multimedia", 
         "El núcleo operativo de la plataforma. Presenta un catálogo asíncrono donde los usuarios pueden aplicar filtros multicriterio (tipo de inmueble, estado legal, rango de precios). "
         "Internamente, este módulo gestiona la carga de archivos binarios (imágenes) validando su extensión y asegurando un almacenamiento local estructurado. "
         "Garantiza que la información expuesta al cliente sea precisa y esté actualizada en tiempo real.", 
         "capturas/catalogo.png"),
        
        ("4.3. Módulo de Propietarios e Interacción UX (Modales Asíncronos)", 
         "Diseñado bajo los principios de Usabilidad (UX). Permite el registro de nuevos dueños de propiedades. "
         "Su innovación radica en la capacidad de registrar un propietario en medio del flujo de creación de un inmueble utilizando una ventana modal, "
         "lo cual evita la ruptura del flujo cognitivo del agente inmobiliario, ahorrando tiempo valioso durante la digitación de datos.", 
         "capturas/crear.png"),
        
        ("4.4. Módulo Gerencial: Dashboard Analítico e Inteligencia de Negocios", 
         "La herramienta para la toma de decisiones. Este panel consolida el volumen total de transacciones y expone KPIs (Key Performance Indicators). "
         "Utilizando la librería Chart.js, transforma registros tabulares en gráficos comprensibles (gráficas de dona y barras), "
         "permitiendo al administrador detectar tendencias de mercado, como por ejemplo, si existe mayor demanda en alquileres frente a ventas en el mes en curso.", 
         "capturas/dashboard.png")
    ]

    for titulo, desc, img_path in modulos:
        doc.add_heading(titulo, level=2)
        doc.add_paragraph(desc)
        try:
            doc.add_picture(img_path, width=Inches(6.0))
            last_paragraph = doc.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"Figura: Vista de interfaz de usuario del {titulo[4:]}.\n")
        except Exception as e:
            doc.add_paragraph(f"[No se pudo cargar la imagen {img_path}]")
        doc.add_page_break()

    # 6. Diagrama E-R
    doc.add_heading('5. Arquitectura del Modelo de Base de Datos Relacional', level=1)
    doc.add_paragraph(
        "El corazón de la persistencia de datos reside en un modelo relacional estrictamente normalizado. "
        "El esquema conceptual, lógico y físico se compone de 8 tablas interconectadas mediante restricciones de clave externa (Foreign Keys). "
        "Se utilizaron restricciones de integridad para asegurar que no existan 'registros huérfanos' (por ejemplo, contratos sin una propiedad asignada). "
        "Las entidades principales abarcan Usuarios, Roles, Propiedades, Propietarios, Visitas, Contratos y tablas paramétricas de Tipos y Estados."
    )
    
    diagram_path = r"c:\Users\wilso\Desktop\emergentes TRA\DIAGRAMA DE BASE DE DATOS .png"
    try:
        doc.add_picture(diagram_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Figura: Diagrama Entidad-Relación (Crow's Foot Notation) generado a partir del esquema estructural.")
    except Exception as e:
        doc.add_paragraph("[Pegar aquí el Diagrama E-R]")

    doc.add_page_break()

    # 7. Tecnologías
    doc.add_heading('6. Descripción técnica de Herramientas y Stack Tecnológico', level=1)
    doc.add_paragraph("El desarrollo del sistema se orquestó combinando tecnologías modernas que garantizan un alto rendimiento y mantenibilidad:")
    doc.add_paragraph(
        "• Lenguaje de Programación: Python 3.13. Seleccionado por su sintaxis legible, vasto ecosistema de paquetes y alto rendimiento en el procesamiento de lógica de negocio en el backend."
    )
    doc.add_paragraph(
        "• Framework Web Backend: Flask. Utilizado en conjunción con el patrón Application Factory. "
        "Proporciona la infraestructura para el enrutamiento HTTP, la gestión del ciclo de vida de la petición/respuesta y la inyección de dependencias."
    )
    doc.add_paragraph(
        "• Mapeo Objeto-Relacional (ORM): Flask-SQLAlchemy. Herramienta vital que traduce los objetos de Python a sentencias SQL nativas. "
        "Su uso mitiga drásticamente las vulnerabilidades de inyección SQL y facilita las futuras migraciones de motores de bases de datos."
    )
    doc.add_paragraph(
        "• Motor de Base de Datos: SQLite3. Empleado por su naturaleza 'serverless' para el ciclo de desarrollo iterativo. "
        "Debido a la abstracción del ORM, el paso a producción con PostgreSQL o MySQL es completamente transparente."
    )
    doc.add_paragraph(
        "• Interfaz de Usuario (Frontend): HTML5 Semántico, CSS3, y Bootstrap 5. El framework Bootstrap se configuró para aprovechar su sistema de grillas (Grid System), "
        "asegurando que el sistema sea 'Responsive' y pueda ser operado cómodamente desde dispositivos móviles (smartphones) por los agentes en campo."
    )
    doc.add_paragraph(
        "• Control de Sesiones y Seguridad: Integración de Flask-Login para el manejo seguro de cookies de sesión HTTP-Only, "
        "junto con Werkzeug.security para la derivación de claves criptográficas."
    )

    doc.add_page_break()

    # 8. Herramientas IA
    doc.add_heading('7. Integración y uso de herramientas de Inteligencia Artificial', level=1)
    doc.add_paragraph(
        "Durante el ciclo de vida del desarrollo del software, se incorporaron Asistentes de Inteligencia Artificial (basados en Modelos de Lenguaje Grande o LLMs) "
        "como un componente integral de la metodología de ingeniería. La adopción de esta tecnología representó un salto cuantitativo en la velocidad de desarrollo."
    )
    doc.add_paragraph(
        "Finalidad de Uso: La IA se empleó primariamente en tres frentes: 1) Generación de código 'boilerplate' (plantillas repetitivas) para los modelos de base de datos. "
        "2) Resolución algorítmica de conflictos en las relaciones de multiplicidad compleja (One-to-Many y Many-to-Many) en SQLAlchemy. "
        "3) Análisis estático de código en tiempo real para predecir vulnerabilidades y optimizar bucles de renderizado en plantillas Jinja2."
    )
    doc.add_paragraph(
        "Reflexión de Impacto: El uso de Inteligencia Artificial no sustituyó el razonamiento analítico del desarrollador, sino que lo potenció. "
        "El desafío se trasladó de 'escribir líneas de código' a 'orquestar la arquitectura'. El desarrollador asumió un rol de Arquitecto y Director, "
        "dictando las reglas de negocio precisas y auditando matemáticamente las soluciones sugeridas por la IA. "
        "Esto permitió construir una aplicación de nivel empresarial en una fracción del tiempo tradicional."
    )

    doc.add_page_break()

    # 9. Conclusiones
    doc.add_heading('8. Conclusiones Finales y Trabajo Futuro', level=1)
    doc.add_heading('8.1. Conclusiones', level=2)
    doc.add_paragraph(
        "El proyecto 'Sistema de Gestión Inmobiliaria TEM742' culminó con éxito rotundo. Se demostró fehacientemente "
        "que la adopción de una arquitectura monolítica con el patrón Application Factory es altamente efectiva para proyectos empresariales "
        "en fase de crecimiento. El sistema cumple estrictamente con los requerimientos operativos planteados. "
        "Las 8 tablas de la base de datos soportan toda la lógica relacional sin presentar cuellos de botella."
    )
    doc.add_paragraph(
        "Asimismo, se constató que la automatización de procesos clave (como el registro de propiedades en catálogo público y el cálculo estadístico en el Dashboard) "
        "reduce el margen de error operativo de una agencia a niveles casi nulos, lo cual se traduce en un mejor retorno de inversión (ROI) a mediano plazo."
    )
    
    doc.add_heading('8.2. Aprendizajes Adquiridos', level=2)
    doc.add_paragraph(
        "El desarrollo exhaustivo de este proyecto consolidó competencias técnicas avanzadas en diversas áreas de la ingeniería de software moderno:"
    )
    doc.add_paragraph(
        "• Ingeniería de Backend: Se dominó la estructuración modular con Blueprints, resolviendo problemas de dependencias circulares comunes en Python."
    )
    doc.add_paragraph(
        "• Modelado de Datos: Se perfeccionó la habilidad de mapear procesos de negocio del mundo real (Dueños, Propiedades, Visitas) hacia estructuras relacionales abstractas utilizando un ORM avanzado."
    )
    doc.add_paragraph(
        "• Experiencia de Usuario (UX): Se comprendió la criticidad de diseñar interfaces no solo bonitas, sino ergonómicas, que eviten la fatiga del usuario final."
    )

    doc.add_heading('8.3. Trabajo Futuro y Escalabilidad', level=2)
    doc.add_paragraph(
        "Como siguientes pasos en la evolución del producto, se proyecta la integración de una API RESTful pública para aplicaciones móviles nativas, "
        "la migración de la base de datos a un clúster de PostgreSQL en la nube (AWS o Render), "
        "y la implementación de pasarelas de pago electrónico para automatizar la reserva de alquileres directamente desde el portal web."
    )

    output_path = 'INFORME_FINAL_TEM742_APA7.docx'
    doc.save(output_path)
    print(f"¡Documento guardado EXITOSAMENTE como '{output_path}'! (Contiene múltiples páginas).")

if __name__ == '__main__':
    create_extended_apa_document()
